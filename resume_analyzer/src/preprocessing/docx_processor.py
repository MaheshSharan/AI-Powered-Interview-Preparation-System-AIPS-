import logging
from typing import Dict, List, Optional, Tuple
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path
import json

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DOCXProcessor:
    def __init__(self):
        """Initialize DOCX processor."""
        # Common resume section headers
        self.section_headers = {
            'contact': ['contact', 'contact information', 'email', 'phone', 'address'],
            'summary': ['summary', 'professional summary', 'objective', 'career objective'],
            'experience': ['experience', 'work experience', 'employment history', 'professional experience'],
            'education': ['education', 'academic background', 'qualifications'],
            'skills': ['skills', 'technical skills', 'core competencies', 'expertise'],
            'projects': ['projects', 'personal projects', 'portfolio']
        }

    def extract_text_with_formatting(self, docx_path: str) -> Tuple[str, Dict]:
        """
        Extract text with formatting preservation using python-docx.
        
        Args:
            docx_path: Path to DOCX file
            
        Returns:
            Tuple of (extracted text, formatting information)
        """
        try:
            doc = Document(docx_path)
            text_content = []
            formatting_info = {
                'paragraphs': [],
                'styles': [],
                'fonts': [],
                'sizes': [],
                'colors': [],
                'alignments': []
            }
            
            # Process each paragraph
            for para in doc.paragraphs:
                if not para.text.strip():
                    continue
                    
                # Extract text
                text_content.append(para.text)
                
                # Extract formatting
                formatting_info['paragraphs'].append({
                    'text': para.text,
                    'style': para.style.name,
                    'level': para.style.paragraph_format.level
                })
                
                # Extract run-level formatting
                for run in para.runs:
                    formatting_info['styles'].append(run.style.name)
                    formatting_info['fonts'].append(run.font.name)
                    formatting_info['sizes'].append(run.font.size.pt if run.font.size else 0)
                    formatting_info['colors'].append(run.font.color.rgb if run.font.color else None)
                    formatting_info['alignments'].append(para.alignment)
            
            return '\n'.join(text_content), formatting_info
            
        except Exception as e:
            logger.error(f"Error extracting text from {docx_path}: {str(e)}")
            raise

    def extract_styles(self, docx_path: str) -> Dict:
        """
        Extract document styles and formatting.
        
        Args:
            docx_path: Path to DOCX file
            
        Returns:
            Dictionary containing style information
        """
        try:
            doc = Document(docx_path)
            styles = {
                'paragraph_styles': [],
                'character_styles': [],
                'table_styles': [],
                'section_properties': []
            }
            
            # Extract paragraph styles
            for style in doc.styles:
                if style.type == 1:  # Paragraph style
                    styles['paragraph_styles'].append({
                        'name': style.name,
                        'base_style': style.base_style.name if style.base_style else None,
                        'font': style.font.name if style.font else None,
                        'size': style.font.size.pt if style.font and style.font.size else None
                    })
                elif style.type == 2:  # Character style
                    styles['character_styles'].append({
                        'name': style.name,
                        'font': style.font.name if style.font else None,
                        'size': style.font.size.pt if style.font and style.font.size else None
                    })
            
            # Extract section properties
            for section in doc.sections:
                styles['section_properties'].append({
                    'page_height': section.page_height.cm,
                    'page_width': section.page_width.cm,
                    'left_margin': section.left_margin.cm,
                    'right_margin': section.right_margin.cm,
                    'top_margin': section.top_margin.cm,
                    'bottom_margin': section.bottom_margin.cm
                })
            
            return styles
            
        except Exception as e:
            logger.error(f"Error extracting styles from {docx_path}: {str(e)}")
            raise

    def identify_sections(self, text: str) -> Dict[str, str]:
        """
        Identify different sections in the resume text.
        
        Args:
            text: Extracted text from DOCX
            
        Returns:
            Dictionary mapping section names to their content
        """
        sections = {}
        lines = text.split('\n')
        current_section = None
        current_content = []

        for line in lines:
            line = line.strip().lower()
            if not line:
                continue

            # Check if line is a section header
            for section, headers in self.section_headers.items():
                if any(header in line for header in headers):
                    if current_section and current_content:
                        sections[current_section] = '\n'.join(current_content)
                    current_section = section
                    current_content = []
                    break
            else:
                if current_section:
                    current_content.append(line)

        # Add the last section
        if current_section and current_content:
            sections[current_section] = '\n'.join(current_content)

        return sections

    def extract_contact_info(self, contact_text: str) -> Dict[str, str]:
        """
        Extract contact information from contact section.
        
        Args:
            contact_text: Text from contact section
            
        Returns:
            Dictionary with contact information
        """
        contact_info = {
            'email': '',
            'phone': '',
            'location': '',
            'linkedin': ''
        }

        lines = contact_text.split('\n')
        for line in lines:
            line = line.lower()
            if '@' in line:
                contact_info['email'] = line.strip()
            elif any(char.isdigit() for char in line):
                contact_info['phone'] = line.strip()
            elif 'linkedin.com' in line:
                contact_info['linkedin'] = line.strip()
            else:
                contact_info['location'] = line.strip()

        return contact_info

    def process_docx(self, docx_path: str) -> Dict:
        """
        Process DOCX file and extract structured information.
        
        Args:
            docx_path: Path to DOCX file
            
        Returns:
            Dictionary containing structured resume information
        """
        try:
            # Extract text with formatting
            text, formatting_info = self.extract_text_with_formatting(docx_path)
            
            # Extract styles
            styles = self.extract_styles(docx_path)
            
            # Identify sections
            sections = self.identify_sections(text)
            
            # Extract contact information
            contact_info = self.extract_contact_info(sections.get('contact', ''))
            
            return {
                'raw_text': text,
                'sections': sections,
                'contact_info': contact_info,
                'formatting_info': formatting_info,
                'styles': styles,
                'metadata': {
                    'filename': Path(docx_path).name,
                    'file_size': Path(docx_path).stat().st_size,
                    'page_count': len(text.split('\n\n'))
                }
            }
            
        except Exception as e:
            logger.error(f"Error processing DOCX {docx_path}: {str(e)}")
            raise 